import os
import logging
from typing import List, Tuple
from PyPDF2 import PdfReader
from docx import Document
import markdown

logger = logging.getLogger(__name__)

class DocumentParser:
    @staticmethod
    def parse(file_path: str, file_type: str) -> Tuple[List[str], dict]:
        """解析文档，返回 (文本段落列表, 解析信息)
        
        解析信息包含:
        - total_pages: 总页数（PDF）
        - total_paragraphs: 总段落数
        - extracted_chars: 提取的字符数
        - quality_warning: 质量警告（如提取内容过少）
        """
        if file_type == "pdf":
            texts, info = DocumentParser._parse_pdf(file_path)
        elif file_type == "docx":
            texts, info = DocumentParser._parse_docx(file_path)
        elif file_type == "txt":
            texts, info = DocumentParser._parse_txt(file_path)
        elif file_type == "md":
            texts, info = DocumentParser._parse_md(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")
        
        # 质量检查
        total_chars = sum(len(t) for t in texts)
        info["extracted_chars"] = total_chars
        
        if total_chars < 100:
            info["quality_warning"] = f"提取内容过少（仅 {total_chars} 字符），可能解析失败"
            logger.warning(f"文档 {file_path} 提取内容过少: {total_chars} 字符")
        elif total_chars < 500:
            info["quality_warning"] = f"提取内容较少（{total_chars} 字符），请检查文档内容"
        else:
            info["quality_warning"] = None
            
        return texts, info

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[List[str], dict]:
        """解析 PDF，支持加密文档和错误处理"""
        info = {"total_pages": 0}
        try:
            reader = PdfReader(file_path)
            info["total_pages"] = len(reader.pages)
            
            # 检查是否加密
            if reader.is_encrypted:
                try:
                    reader.decrypt("")  # 尝试空密码
                except Exception:
                    raise ValueError("PDF 文件已加密，无法解析")
            
            texts = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        texts.append(text.strip())
                except Exception as e:
                    logger.warning(f"PDF 第 {i+1} 页解析失败: {e}")
                    continue
                    
            return texts, info
            
        except Exception as e:
            logger.error(f"PDF 解析失败: {e}")
            raise ValueError(f"PDF 解析失败: {str(e)}")

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[List[str], dict]:
        """解析 DOCX"""
        try:
            doc = Document(file_path)
            texts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    texts.append(para.text.strip())
            return texts, {"total_paragraphs": len(doc.paragraphs)}
        except Exception as e:
            logger.error(f"DOCX 解析失败: {e}")
            raise ValueError(f"DOCX 解析失败: {str(e)}")

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[List[str], dict]:
        """解析 TXT，支持多种编码"""
        encodings = ["utf-8", "gbk", "gb2312", "utf-16"]
        content = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as f:
                    content = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if content is None:
            raise ValueError("无法识别文件编码，请使用 UTF-8 编码")
        
        # 按段落分割（双换行）
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        return paragraphs, {"encoding": used_encoding}

    @staticmethod
    def _parse_md(file_path: str) -> Tuple[List[str], dict]:
        """解析 Markdown"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                md_content = f.read()
            # 将 Markdown 转为纯文本
            html = markdown.markdown(md_content)
            # 去除 HTML 标签
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "html.parser")
            text = soup.get_text()
            paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
            return paragraphs, {"total_paragraphs": len(paragraphs)}
        except Exception as e:
            logger.error(f"Markdown 解析失败: {e}")
            raise ValueError(f"Markdown 解析失败: {str(e)}")